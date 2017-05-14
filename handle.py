from docx import Document
import os
import shelve
import re
import pprint


def doc2data2017():
    docs = os.listdir('download/')
    meatdocs = []
    for doc in docs:
        if doc.startswith('~'):
            docs.remove(doc)
        if '猪肉' in doc and '蔬菜' not in doc:
            meatdocs.append(doc)
            docs.remove(doc)

    for doc in meatdocs:
        os.remove('download/' + doc)

    docs.sort(reverse=True)
    data = []
    for path in docs[4:13]:
        subdata = []
        doc = Document('download/' + path)
        for row in doc.tables[2].rows:
            cells = []
            for cell in row.cells:
                if cell.text == '序号' or cell.text == '日期' or cell.text == '名称':
                    break
                # print(cell.text, end=' ')
                cells.append(cell.text)

            if cells != []:
                subdata.append(tuple(cells))
            # print('')
        data.append(subdata)
    # print(data)
    datafile = shelve.open('data')
    datafile['2017'] = data
    datafile.close()


def doc2data_mo_se():
    paths = []
    for path in os.listdir('download/'):
        if path.endswith('安全报告.docx') and not path.startswith('~'):
            paths.append(path)

    moRegex = re.compile(r'(\d+)年(\d+)月份')
    seRegex = re.compile(r'(\d+)年第(\S+)季度')
    mopaths = []
    sepaths = []
    for path in paths:
        month = moRegex.search(path)
        if month is not None:
            mopaths.append((month.group(), path))
        else:
            season = seRegex.search(path)
            if season is not None:
                sepaths.append((season.group(), path))

    # 蔬菜种类按月统计
    moDatas = {}
    for month, path in mopaths:
        doc = Document('download/' + path)
        moData = {}
        for row in doc.tables[6].rows[1:-1]:
            cells = row.cells
            key = cells[0].text
            if key == '芸薹属类':
                key = '芸苔属类'
            elif key == '其他':
                key = '其他蔬菜品种'
            moData[key] = (int(cells[1].text), int(cells[2].text))
        moDatas[month] = moData

    # 蔬菜种类按季度统计
    seDatas = {}
    for season, path in sepaths:
        se = seRegex.search(season)
        if int(se.group(1)) > 2015 or (int(se.group(1)) == 2015 and se.group(2) == '四'):
            doc = Document('download/' + path)
            seData = {}
            for row in doc.tables[6].rows[1:-1]:
                cells = row.cells
                key = cells[0].text
                if key == '芸薹属类':
                    key = '芸苔属类'
                elif key == '其他':
                    key = '其他蔬菜品种'
                seData[key] = (int(cells[1].text), int(cells[2].text))
            if se.group(2) == '一':
                see = 1
            elif se.group(2) == '二':
                see = 2
            elif se.group(2) == '三':
                see = 3
            else:
                see = 4
            seDatas[se.group(1) + '年第' + str(see) + '季度'] = seData

    # 将按月统计数据转化为按季度统计数据
    for month, datas in moDatas.items():
        mo = moRegex.search(month)
        year = int(mo.group(1))
        month = int(mo.group(2))
        if month <= 3:
            see = 1
        elif month <= 6:
            see = 2
        elif month <= 9:
            see = 3
        else:
            see = 4
        default = {'其他蔬菜品种': (0, 0),
                   '叶菜类': (0, 0),
                   '根茎类': (0, 0),
                   '水果': (0, 0),
                   '水生类': (0, 0),
                   '瓜类': (0, 0),
                   '芸苔属类': (0, 0),
                   '茄果类': (0, 0),
                   '豆类': (0, 0),
                   '食用菌': (0, 0),
                   '鳞茎类': (0, 0)}
        odatas = seDatas.get(str(year) + '年第' + str(see) + '季度', default)
        for k, (v1, v2) in datas.items():
            odatas[k] = (odatas[k][0] + v1, odatas[k][1] + v2)
        seDatas[str(year) + '年第' + str(see) + '季度'] = odatas
    # 保存蔬菜种类数据
    datafile = shelve.open('data')
    datafile['vegtype_season'] = seDatas

    # 渠道比例按月统计
    moDatas_chan = {}
    for month, path in mopaths:
        doc = Document('download/' + path)
        subdata = {}
        for row in doc.tables[1].rows[1:5]:
            cells = row.cells[1:-1]
            key = cells[0].text
            if cells[0].text == '超市':
                key = '商场超市'
            subdata[key] = (int(cells[1].text), int(cells[2].text))
        moDatas_chan[month] = subdata

    # 渠道比例按季度统计
    seDatas_chan = {}
    for season, path in sepaths:
        doc = Document('download/' + path)
        subdata = {}
        for row in doc.tables[1].rows[1:5]:
            cells = row.cells[1:-1]
            key = cells[0].text
            if cells[0].text == '超市':
                key = '商场超市'
            subdata[key] = (int(cells[1].text), int(cells[2].text))
        seDatas_chan[season] = subdata

    # 渠道比例从按月统计转换为按季度统计
    for month, datas in moDatas_chan.items():
        mo = moRegex.search(month)
        year = int(mo.group(1))
        month = int(mo.group(2))
        if month <= 3:
            see = '一'
        elif month <= 6:
            see = '二'
        elif month <= 9:
            see = '三'
        else:
            see = '四'
        default = {'生产基地': (0, 0),
                   '农贸市场': (0, 0),
                   '商场超市': (0, 0),
                   '批发市场': (0, 0)}
        odatas = seDatas_chan.get(str(year) + '年第' + see + '季度', default)
        for k, (v1, v2) in datas.items():
            odatas[k] = (odatas[k][0] + v1, odatas[k][1] + v2)
        seDatas_chan[str(year) + '年第' + str(see) + '季度'] = odatas

    # 保存渠道比例数据
    datafile['chantype_season'] = seDatas_chan
    datafile.close()
    return seDatas_chan


se = doc2data_mo_se()
pprint.pprint(se)
